import docx
import os

os.chdir(r"C:\Python\automate_the_boring_stuff\Chapter 13 - Working with Excelsheets")
d = docx.Document()
d.add_paragraph("hello penisholz")
d.add_paragraph("anotherone cunt")

p = d.paragraphs[0]

p.add_run("this is a new run")
p.runs[1].bold = True



d.save("demo2.docx")