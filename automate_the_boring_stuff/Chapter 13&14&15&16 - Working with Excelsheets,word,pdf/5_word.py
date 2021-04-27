import docx
import os

os.chdir(r"C:\Python\automate_the_boring_stuff\Chapter 13 - Working with Excelsheets")
d = docx.Document("demo.docx")

print(d.paragraphs[0].text)
print(d.paragraphs[1].text)

p = d.paragraphs[1]
print(p.runs[1].text)

p.runs[3].underline = True
p.runs[3].text = "italic and underlined."

p.style = "Title" # Style geändert

d.save("new.docx")

d = docx.Document

